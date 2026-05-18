"""RAG 文档处理模块

支持多格式文档加载和智能分块:
- Markdown / TXT / PDF / Word / HTML 等格式
- 基于 Markdown 标题层次的结构感知分块
- Token 感知的智能分块
"""
from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional


class Document:
    """文档对象"""

    def __init__(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
        doc_id: Optional[str] = None,
    ):
        self.content = content
        self.metadata = metadata or {}
        self.doc_id = doc_id or self._generate_id()

    def _generate_id(self) -> str:
        import hashlib
        return hashlib.md5(self.content.encode()).hexdigest()[:16]

    def __repr__(self) -> str:
        title = self.metadata.get("title", self.doc_id[:16])
        return f"Document(id={self.doc_id}, title={title}, length={len(self.content)})"


class DocumentProcessor:
    """文档处理器 - 多格式解析 + 智能分块"""

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ========== 文档加载 ==========

    def load_file(self, file_path: str) -> Optional[Document]:
        """从文件加载文档

        支持格式: .txt, .md, .py, .json, .html, .pdf, .docx
        """
        if not os.path.exists(file_path):
            print(f"❌ 文件不存在: {file_path}")
            return None

        ext = os.path.splitext(file_path)[1].lower()

        loaders = {
            ".txt": self._load_text,
            ".md": self._load_text,
            ".py": self._load_text,
            ".js": self._load_text,
            ".ts": self._load_text,
            ".json": self._load_text,
            ".html": self._load_text,
            ".csv": self._load_text,
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
        }

        loader = loaders.get(ext, self._load_text)
        try:
            content = loader(file_path)
            if not content or not content.strip():
                print(f"⚠️ 文件内容为空: {file_path}")
                return None

            return Document(
                content=content,
                metadata={
                    "source": file_path,
                    "filename": os.path.basename(file_path),
                    "extension": ext,
                },
            )
        except Exception as e:
            print(f"❌ 文件加载失败 {file_path}: {e}")
            return None

    def load_text(self, text: str, doc_id: Optional[str] = None, metadata: Optional[Dict] = None) -> Document:
        """从文本字符串加载文档"""
        return Document(content=text, metadata=metadata or {}, doc_id=doc_id)

    def _load_text(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def _load_pdf(self, file_path: str) -> str:
        """加载 PDF 文件（使用 PyPDF2 或 pdfplumber）"""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n\n".join(text_parts)
        except ImportError:
            pass

        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file_path)
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            pass

        raise ImportError("需要安装 pdfplumber 或 PyPDF2: pip install pdfplumber")

    def _load_docx(self, file_path: str) -> str:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs)
        except ImportError:
            raise ImportError("需要安装 python-docx: pip install python-docx")

    # ========== 智能分块 ==========

    def split_document(self, document: Document) -> List[Dict[str, Any]]:
        """将文档分割为适合检索的文本块

        策略：
        1. 对于 Markdown 文本，按标题层次分割
        2. 按段落分割
        3. 基于 Token 数量合并为合适大小的块
        """
        text = document.content

        # 检测是否为 Markdown
        if self._is_markdown(text):
            paragraphs = self._split_by_headings(text)
        else:
            paragraphs = self._split_by_paragraphs(text)

        return self._merge_paragraphs_to_chunks(paragraphs, document.metadata)

    def _is_markdown(self, text: str) -> bool:
        """检测是否为 Markdown 格式"""
        return bool(re.search(r'^#{1,6}\s', text, re.MULTILINE))

    def _split_by_headings(self, text: str) -> List[Dict[str, Any]]:
        """按 Markdown 标题层次分割"""
        lines = text.splitlines()
        heading_stack = []
        paragraphs = []
        buf = []
        char_pos = 0

        for ln in lines:
            raw = ln
            # 处理标题行
            stripped = raw.strip()
            if stripped.startswith("#"):
                if buf:
                    content = "\n".join(buf).strip()
                    if content:
                        paragraphs.append({
                            "content": content,
                            "heading_path": " > ".join(heading_stack) if heading_stack else None,
                        })
                    buf = []

                level = len(raw) - len(raw.lstrip('#'))
                title = stripped.lstrip('#').strip()
                if level <= 0:
                    level = 1
                if level <= len(heading_stack):
                    heading_stack = heading_stack[:level - 1]
                heading_stack.append(title)
                char_pos += len(raw) + 1
                continue

            if stripped == "" and buf:
                content = "\n".join(buf).strip()
                if content:
                    paragraphs.append({
                        "content": content,
                        "heading_path": " > ".join(heading_stack) if heading_stack else None,
                    })
                buf = []
            else:
                buf.append(raw)
            char_pos += len(raw) + 1

        # 处理最后的缓冲区
        if buf:
            content = "\n".join(buf).strip()
            if content:
                paragraphs.append({
                    "content": content,
                    "heading_path": " > ".join(heading_stack) if heading_stack else None,
                })

        if not paragraphs:
            paragraphs = [{"content": text, "heading_path": None}]

        return paragraphs

    def _split_by_paragraphs(self, text: str) -> List[Dict[str, Any]]:
        """按段落分割"""
        # 按空行分割
        parts = re.split(r'\n\s*\n', text)
        return [
            {"content": p.strip(), "heading_path": None}
            for p in parts if p.strip()
        ] or [{"content": text, "heading_path": None}]

    def _merge_paragraphs_to_chunks(
        self, paragraphs: List[Dict[str, Any]], metadata: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """基于 Token 数量将段落合并为合适大小的块"""
        chunks = []
        current_chunks = []
        current_tokens = 0

        for para in paragraphs:
            para_tokens = self._estimate_tokens(para["content"])

            if current_tokens + para_tokens > self.chunk_size and current_chunks:
                # 生成当前块
                content = "\n\n".join(c["content"] for c in current_chunks)
                heading = next(
                    (c["heading_path"] for c in reversed(current_chunks) if c.get("heading_path")),
                    None,
                )
                chunks.append({
                    "content": content,
                    "heading_path": heading,
                    "metadata": metadata,
                    "char_count": len(content),
                })

                # 重叠: 保留最后几个段落
                if self.chunk_overlap > 0:
                    kept = []
                    kept_tokens = 0
                    for c in reversed(current_chunks):
                        t = self._estimate_tokens(c["content"])
                        if kept_tokens + t > self.chunk_overlap:
                            break
                        kept.append(c)
                        kept_tokens += t
                    current_chunks = list(reversed(kept))
                    current_tokens = kept_tokens
                else:
                    current_chunks = []
                    current_tokens = 0

            current_chunks.append(para)
            current_tokens += para_tokens

        # 最后一个块
        if current_chunks:
            content = "\n\n".join(c["content"] for c in current_chunks)
            heading = next(
                (c["heading_path"] for c in reversed(current_chunks) if c.get("heading_path")),
                None,
            )
            chunks.append({
                "content": content,
                "heading_path": heading,
                "metadata": metadata,
                "char_count": len(content),
            })

        return chunks

    def _estimate_tokens(self, text: str) -> int:
        """估算 Token 数量（支持中英文混合）"""
        cjk = sum(1 for ch in text if self._is_cjk(ch))
        other_tokens = len([t for t in text.split() if t])
        return cjk + other_tokens

    def _is_cjk(self, ch: str) -> bool:
        """判断是否为 CJK 字符"""
        code = ord(ch)
        return (
            0x4E00 <= code <= 0x9FFF
            or 0x3400 <= code <= 0x4DBF
            or 0x20000 <= code <= 0x2A6DF
            or 0xF900 <= code <= 0xFAFF
        )
